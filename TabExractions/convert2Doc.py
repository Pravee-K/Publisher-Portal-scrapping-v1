import json
from docx import Document

# Load JSON data from a file
with open(r'C:\Users\Admin\Downloads\Publisher-Portal-scrapping\Publisher-Portal-scrapping-c3f406f7401c74c8741a31781b0a23a10a2fcf9f\TabExractions\TabSupport\data\ClgInfoOutput.json', 'r') as json_file:
    data = json.load(json_file)

# Create a new Document
doc = Document()

# Add a title to the Word document
doc.add_heading('JSON Data to Word Conversion', 0)

# Recursively add JSON content to the document
def add_content_to_doc(data, doc, level=1):
    if isinstance(data, dict):
        for key, value in data.items():
            doc.add_heading(f"{key}:", level=level)
            add_content_to_doc(value, doc, level+1)
    elif isinstance(data, list):
        for index, item in enumerate(data):
            doc.add_heading(f"Item {index+1}:", level=level)
            add_content_to_doc(item, doc, level+1)
    else:
        doc.add_paragraph(str(data))

# Add JSON data to the document
add_content_to_doc(data, doc)

# Save the document as a Word file
doc.save('output.docx')

print("JSON data has been successfully converted to a Word document.")
